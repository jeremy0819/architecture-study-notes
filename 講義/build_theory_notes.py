# -*- coding: utf-8 -*-
"""產生「敷地計畫理論名著精要（蒸餾・可寫句）」Word 檔。

蒸餾三本經典的核心方法論，接上台灣法規與審議規範，寫成考場可默寫的短句。
與 references/理論名著精要.md 同步（該 MD 為單一來源）。
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK = 'Microsoft JhengHei'
ACCENT = RGBColor(0x1F, 0x4E, 0x79)

doc = Document()

def style_font(style_name, size=None, bold=None, color=None):
    st = doc.styles[style_name]
    st.font.name = CJK
    rpr = st.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(a), CJK)
    if size is not None: st.font.size = Pt(size)
    if bold is not None: st.font.bold = bold
    if color is not None: st.font.color.rgb = color

style_font('Normal', size=10.5)
for h, sz in (('Heading 1', 16), ('Heading 2', 13), ('Heading 3', 11.5)):
    style_font(h, size=sz, bold=True, color=ACCENT)
style_font('Title', size=26, bold=True, color=ACCENT)

def run_font(run):
    run.font.name = CJK
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(a), CJK)

def P(text='', size=10.5, bold=False, color=None, align=None, space_after=4):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold
    if color: r.font.color.rgb = color
    run_font(r)
    return p

def bullets(items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(it); r.font.size = Pt(size); run_font(r)

def H1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: run_font(r)
    return p

def H2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: run_font(r)
    return p

def H3(t):
    p = doc.add_heading(t, level=3)
    for r in p.runs: run_font(r)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexcolor)
    tcPr.append(sh)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]; r = p.add_run(h)
        r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run_font(r)
        shade(hdr[i], '1F4E79')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]; r = p.add_run(str(val))
            r.font.size = Pt(9); run_font(r)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def box(text, fill='FFF2CC', bold=True, size=10):
    t = doc.add_table(rows=1, cols=1); t.style='Table Grid'
    c = t.rows[0].cells[0]; shade(c, fill); c.text=''
    p = c.paragraphs[0]; r = p.add_run(text); r.font.size=Pt(size); r.font.bold=bold; run_font(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def write_line(label, text):
    """可寫句：粗體標籤＋內容"""
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label); r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = ACCENT; run_font(r1)
    r2 = p.add_run(text); r2.font.size = Pt(10); run_font(r2)

def pagebreak():
    doc.add_page_break()

# ================= 封面 =================
doc.add_paragraph().paragraph_format.space_after = Pt(60)
ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run('敷地計畫\n理論名著精要'); r.font.size = Pt(32); r.font.bold = True; r.font.color.rgb = ACCENT; run_font(r)
P('蒸餾・可寫句｜敷地計畫與都市設計（80150）', size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
P('Lynch & Hack《敷地計畫》・Edward White《基地分析》・詹氏《圖解都市計劃》', size=11,
  align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x40,0x40,0x40))
doc.add_paragraph().paragraph_format.space_after = Pt(24)
box('把三本經典的核心方法論，蒸餾成考場「寫得完、讀得懂、不晦澀」的短句，並接上台灣法規與審議規範——'
    '讓申論從「背料庫」升級為「有理論靠山的論述」。', fill='DEEBF7', bold=False)
box('⚠️ 可信度：本檔為上述作者「公認核心方法論」之蒸餾與轉譯，非逐字引用您手上之特定版本。'
    '理論精神可信，惟具體頁碼字句請以原著為準；台灣法規以全國法規資料庫最新版為準。')
pagebreak()

# ================= 0 定位 =================
H1('0. 一句話定位：三本書各補你哪一塊')
table(['書', '它最強的一件事', '補你的弱點', '對應考題'], [
    ['Lynch & Hack《敷地計畫》', '為什麼這樣配置（行為、適配、感受）', '只會排機能、講不出為何這樣排', '108、110、114'],
    ['Edward White《基地分析》', '怎麼分析＋怎麼畫成圖（系統清單＋圖解）', '分析漏項、圖畫不出、圖文不對應', '110、112、每年設計題'],
    ['詹氏《圖解都市計劃》', '法定框架（都計三部＋審議）', '法源講不清、審議接不上', '109、都審類、每年設計說明'],
], widths=[4.5, 5.5, 4.5, 3])
box('記法：White 教你「讀懂基地並畫出來」→ Lynch 教你「為何而設計」→ 圖解都計 教你「合不合法、過不過審」。'
    '三者串起來，就是一份完整的敷地計畫論述。', fill='E2EFDA', bold=False)
pagebreak()

# ================= 1 Lynch =================
H1('1. Lynch & Hack《敷地計畫》第三版——蒸餾五要點')
P('這本是本科目的「聖經」，108、114 申論直接點名 Lynch。抓住下面五個觀念，就能為任何題目補上「理論深度」那一層。', bold=True)

H2('1.1 敷地計畫是什麼（開場定義，萬用）')
write_line('蒸餾：', '敷地計畫是「安排戶外實質環境以支持人的行為」，重點不只在擺建築物，更在塑造建築物「之間」的空間。')
write_line('可寫句：', '「敷地計畫的核心不在配置量體，而在塑造量體之間的開放空間，使人的行為得以發生——此即 Lynch 所謂『設計建築物之間的空間』。」')
write_line('接法規／審議：', '呼應都市設計審議「建立整體空間架構、改善環境景觀」之宗旨。')

H2('1.2 空間計畫不是既定條件（programming 觀，108 核心）')
write_line('蒸餾：', '業主給的空間計畫只是「假設」，要放回基地檢討、修正；並分析「新機能置入後，對既有使用者行為與環境的後續效應」。')
write_line('可寫句：', '「依 Lynch 之見，業主預擬之空間計畫不應視為既定條件，而應放回基地環境檢討——先評估新量體置入後對現況使用者行為與環境的後續效應，再據以修正計畫。」')
write_line('接法規／審議：', '對應都市計畫「先調查分析、後擬定計畫」之程序精神。')

H2('1.3 場所＝活動＋地點（behavior setting，配置的底層邏輯）')
write_line('蒸餾：', '空間的意義來自「活動 × 場所」的結合；配置前先設想使用者的行為劇本，再給它對應的場所。')
write_line('可寫句：', '「配置係以使用者行為劇本為本：先設想『誰、在何時、做什麼』，再賦予對應之場所與尺度，使每一處開放空間都有明確的活動承載。」')
write_line('接法規／審議：', '呼應無障礙與通用設計對「使用者友善」之要求。')

H2('1.4 好敷地＝形式與使用的「適配」＋未來彈性（fit & adaptability）')
write_line('蒸餾：', '好的敷地計畫是形式與使用之間的「適配」，並預留未來調整的彈性——不是一次到位的完美，而是能長期演化。')
write_line('可寫句：', '「良好之敷地計畫講求形式與使用之『適配』，並以大跨距、可變隔間等手法預留未來彈性，使空間得隨機能演化而調整（平災轉換亦屬其一）。」')
write_line('接法規／審議：', '接韌性城市之「彈性空間」與公設多目標使用辦法。')

H2('1.5 感受性的基地（sensed site／五元素）')
write_line('蒸餾：', '基地不只是數據，更是被經驗的；意象五元素（路徑、邊界、節點、地標、區域）是描述「人如何感知空間」的語言，可平面用、也可立體化用（114）。')
write_line('可寫句：', '「除量化條件外，應分析基地的『感受性』——以路徑、邊界、節點、地標、區域五元素組織空間經驗，並於立體都市中向垂直延伸（空橋為立體路徑、挑空大廳為立體節點）。」')
write_line('接法規／審議：', '都審對「都市紋理、天際線、視覺廊道」之管制即感受性之落實。')
box('Lynch 一句話總結（考前默念）：先問人怎麼用、怎麼感受，再決定怎麼配置與量體——這就是敷地計畫與「純建築設計」的分野。', fill='E2EFDA', bold=False)
pagebreak()

# ================= 2 White =================
H1('2. Edward T. White《基地分析／序列系統》——蒸餾三要點')
P('White 最有價值的，是把「基地分析」變成不會漏項的系統清單，並主張每一項分析都畫成一張圖。'
  '正好補台灣考生兩個致命傷：分析漏項、圖文不對應。', bold=True)

H2('2.1 基地分析的系統清單（十一類，對照台灣資料來源）')
table(['White 分析類別', '白話', '台灣資料來源（接 110／112）'], [
    ['① 區位 Location', '在區域／都市中的位置', '都計圖、航照、地圖'],
    ['② 鄰里涵構 Neighborhood', '周邊使用、現況、未來', '現勘、都計分區圖'],
    ['③ 大小與分區 Size & Zoning', '尺寸、建蔽容積、退縮、高度', '土地使用分區證明、都審原則'],
    ['④ 法令 Legal', '產權、地役、限制', '地籍、建管指定建築線'],
    ['⑤ 自然實質 Natural', '地形、排水、樹木、土壤、水文', '地形圖、地質所、淹水潛勢圖'],
    ['⑥ 人為實質 Man-made', '既有建物、圍牆、鋪面', '現勘、測量'],
    ['⑦ 交通 Circulation', '人車動線、尖離峰、停留', '交通局流量、現場觀察'],
    ['⑧ 公用設施 Utilities', '水電瓦斯污水、位置容量', '自來水／下水道／台電管線圖'],
    ['⑨ 感官 Sensory', '視覺、噪音、氣味（內外互看）', '現勘攝影、環保監測'],
    ['⑩ 人文 Human & Cultural', '周邊人口、活動、歷史記憶', '統計、口述歷史、文化局（接 112）'],
    ['⑪ 氣候 Climate', '日照、溫度、雨、風、濕度', '中央氣象署'],
], widths=[4.5, 5, 6.5])
write_line('可寫句（列舉題直接用）：', '「基地分析宜以系統化清單全面掃描——涵蓋區位、鄰里涵構、分區法令、自然與人為實質、交通、公用設施、感官、人文與氣候各面向，避免遺漏。」')

H2('2.2 每一項分析都畫成「一張圖」（圖解分析法，本科最大加分點）')
write_line('蒸餾：', 'White 的招牌——用一張圖回答一個分析問題，不要用一段文字。基地分析應是一疊「圖示卡」，不是報告書。')
write_line('可寫句：', '「基地分析採圖解方式呈現：每一分析項目以一張標註清楚之示意圖表達（如風花圖、日照角度、視覺廊道、噪音來向），使『圖說話』——此亦本科目圖文並陳之要求。」')
write_line('接考試：', '這就是評審要看的「版面份量」；把 White 的圖解法內化，就不會再有大片空白。')

H2('2.3 分析是為了導出概念（Analysis → Issue → Concept）')
write_line('蒸餾：', '分析不是中性堆資料，而是為了設計；每一項分析都要指向一個課題（Issue）、再導出一個對策／概念（Concept）。')
write_line('可寫句：', '「分析、課題、對策三者環環相扣：每一項基地條件都應轉為一則課題，並導出對應之設計對策，最終收斂為一句設計主軸。」')
write_line('接法規／審議：', '都審書件的邏輯即「現況分析→課題→設計準則回應」，與此完全同構。')
box('White 一句話總結：把基地讀成一疊會說話的分析圖，讓每張圖都指向一個對策——分析到概念，一路都用畫的。', fill='E2EFDA', bold=False)
pagebreak()

# ================= 3 圖解都計 =================
H1('3. 詹氏《圖解都市計劃》——把法定框架講清楚')
P('這本補你的「法源與審議」底盤。記住下面的體系，設計說明的起手式就有骨架。', bold=True)

H2('3.1 都市計畫三部曲（體系記法）')
write_line('蒸餾：', '主要計畫（指導）→ 細部計畫（管制平台）→ 都市設計（落實界面）。主要計畫定方向與公共設施系統；細部計畫定建蔽容積退縮等可操作管制；都市設計把管制轉譯成實質空間品質。')
write_line('可寫句（起手式升級版）：', '「本案依都市計畫法主要計畫之指導、細部計畫之土地使用管制，並經都市設計審議落實為實質空間品質，形成『指導—管制—落實』之三層架構。」')

H2('3.2 土地使用分區管制＝敷地計畫的「遊戲規則」')
write_line('蒸餾：', '建蔽率、容積率、退縮、高度、停車，都是分區管制給的邊界條件；敷地計畫是在這些邊界內求最佳解。')
write_line('可寫句：', '「量體係於分區管制之邊界條件內求解——建築面積＝基地×建蔽率、容積樓地板＝基地×容積率，並依退縮與高度規定塑造天際線。」')

H2('3.3 都市設計審議＝公共性的守門員')
write_line('蒸餾：', '審議在意的不是單棟美醜，而是對公共領域的貢獻——開放空間可及性、街道介面、行人環境、天際線、綠化。')
write_line('可寫句：', '「本案於都市設計審議面向，著重公共領域之貢獻：留設可及之開放空間、活化街道介面、延續紋理與天際線，體現開放空間之公共性。」')
pagebreak()

# ================= 4 合流表 =================
H1('4. 三書合流：理論 → 法規／審議 → 可寫句 總表')
P('考前掃這張表：看到題眼，反射出「哪個理論＋哪部法＋哪句話」。', bold=True)
table(['題眼', '理論靠山', '接哪部法／審議', '考場可寫句（濃縮）'], [
    ['programming／空間計畫', 'Lynch：計畫非既定', '都市計畫法（先調查後計畫）', '放回基地檢討 programming，評估後續效應再修正。'],
    ['基地分析（列舉）', 'White：十一類清單', '分區管制、建築線', '以系統清單掃描區位至氣候各面向，避免漏項。'],
    ['圖文並陳／版面', 'White：一項一圖', '（考試要求）', '每一分析項目以一張標註圖示表達，讓圖說話。'],
    ['分析→設計', 'White：Issue→Concept', '都審書件邏輯', '每項條件轉為課題、導出對策，收斂為主軸。'],
    ['開放空間公共性', 'Lynch：建物之間的空間', '都市設計審議宗旨', '重點在塑造量體之間的空間與其公共性。'],
    ['使用者友善', 'Lynch：behavior setting', '無障礙設計規範', '以行為劇本為本配置場所，通用設計串聯。'],
    ['彈性／平災轉換', 'Lynch：fit & adaptability', '公設多目標使用辦法', '以適配與彈性預留未來，平時活動、災時轉換。'],
    ['五元素／立體化', 'Lynch：意象五元素', '都審天際線／視覺廊道', '以五元素組織空間經驗，並向垂直立體延伸。'],
    ['法定體系／設計說明', '圖解都計：三部曲', '都市計畫法 §15/§22', '指導—管制—落實三層架構。'],
], widths=[3.5, 4, 4, 5.5])
pagebreak()

# ================= 5 升級句 =================
H1('5. 申論穿插「理論升級句」')
P('法規金句給你「法源句」；這裡給你「理論句」。一段論述裡，理論句放前、法源句接後，份量立刻不同。', bold=True)
table(['原本只寫（法源句）', '升級為（理論句＋法源句）'], [
    ['依都市計畫法配置開放空間。',
     '敷地計畫的重點在塑造量體之間的開放空間（Lynch）；本案依都市計畫法與都審精神，留設可及且具公共性之開放空間。'],
    ['依水利法出流管制設滯洪。',
     '基地應如海綿般留滯滲用（White 感官／氣候分析導出之對策）；依水利法出流管制，設雨花園與地下貯留，逕流不直排。'],
    ['基地分析包含氣候、水文…',
     '基地分析宜以系統清單全面掃描（White），並將每項轉為分析圖與對策；據此導出設計主軸。'],
    ['屋頂設太陽光電。',
     '面對氣候面向之分析（White 氣候類），並回應韌性之替代能源立體化；依再生能源發展條例設屋頂光電。'],
], widths=[6, 11])
box('穿插節奏：一段 ＝ 理論句（為什麼）＋ 設計作為（做什麼）＋ 法源句（有依據）。三句一組，最耐讀也最有份量。',
    fill='FFF2CC')

H1('6. 三本書的「一頁帶走」（背這個就好）')
bullets([
    'Lynch：先問人怎麼用、怎麼感受，再決定配置與量體；空間計畫要放回基地檢討；設計的是建物「之間」的空間。',
    'White：用十一類清單讀基地、每項畫一張圖、每張圖指向一個對策。',
    '圖解都計：主要計畫指導→細部計畫管制→都市設計落實；審議守的是公共性。',
])
box('三句話串起來就是一份敷地計畫的完整心法：讀懂基地（White）→ 為人而設計（Lynch）→ 合法且過審（圖解都計）。',
    fill='DEEBF7', bold=True)

out = '/home/user/architecture-study-notes/講義/敷地計畫理論名著精要.docx'
doc.save(out)
print('saved:', out)
