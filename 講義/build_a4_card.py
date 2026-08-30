# -*- coding: utf-8 -*-
"""敷地計畫與都市設計｜一頁 A4 考場作戰卡。
版面依「作答時間軸」分帶（非模仿圖紙三欄）：頁首開卷帶 → 左欄申論期 → 右欄設計期 → 頁尾交卷帶。
內容經 6 鏡頭策展 + 3 鏡頭對抗投票 + 完整性批判後合併去重。
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK='Microsoft JhengHei'
NAVY=RGBColor(0x1F,0x3B,0x5C); RED=RGBColor(0xC0,0x00,0x00); GRAY=RGBColor(0x44,0x44,0x44)
BODY=9.0; HEAD=10.0

doc=Document(); s=doc.sections[0]
s.page_width=Cm(21.0); s.page_height=Cm(29.7)
s.left_margin=s.right_margin=Cm(0.75); s.top_margin=s.bottom_margin=Cm(0.6)

st=doc.styles['Normal']; st.font.name=CJK; st.font.size=Pt(BODY)
rf=st.element.get_or_add_rPr().get_or_add_rFonts()
for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rf.set(qn(a),CJK)

def _f(r):
    r.font.name=CJK
    rr=r._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rr.set(qn(a),CJK)

def shade(cell,hx):
    tc=cell._tc.get_or_add_tcPr(); e=OxmlElement('w:shd')
    e.set(qn('w:val'),'clear'); e.set(qn('w:fill'),hx); tc.append(e)

def tight(p,before=0,after=0,line=0.92):
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after)
    pf.line_spacing=line

def para(txt,size=BODY,bold=False,color=None,indent=0,before=0,after=0.4):
    p=doc.add_paragraph(); tight(p,before,after)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    r=p.add_run(txt); r.font.size=Pt(size); r.font.bold=bold
    if color: r.font.color.rgb=color
    _f(r); return p

def rich(segs,size=BODY,indent=0,before=0,after=0.4):
    """segs = [(text,bold,color)]"""
    p=doc.add_paragraph(); tight(p,before,after)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    for t,b,c in segs:
        r=p.add_run(t); r.font.size=Pt(size); r.font.bold=b
        if c: r.font.color.rgb=c
        _f(r)
    return p

def head(txt,fill='1F3B5C'):
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'
    c=t.rows[0].cells[0]; shade(c,fill); c.text=''
    c.vertical_alignment=1
    for m in ('top','bottom'):
        c._tc.get_or_add_tcPr()
    p=c.paragraphs[0]; tight(p,0,0,0.9)
    r=p.add_run(txt); r.font.size=Pt(HEAD); r.font.bold=True
    r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _f(r)
    return t

def band(txt,fill='FFF2CC',size=BODY,bold=False):
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'
    c=t.rows[0].cells[0]; shade(c,fill); c.text=''
    p=c.paragraphs[0]; tight(p,0,0,0.9)
    r=p.add_run(txt); r.font.size=Pt(size); r.font.bold=bold; _f(r)
    return t

def cols(n):
    sp=doc.sections[-1]._sectPr
    cs=sp.xpath('./w:cols')[0]
    cs.set(qn('w:num'),str(n)); cs.set(qn('w:space'),'200')
    if n>1: cs.set(qn('w:sep'),'1')

def newsec(numcols):
    from docx.enum.section import WD_SECTION
    ns=doc.add_section(WD_SECTION.CONTINUOUS)
    ns.page_width=Cm(21.0); ns.page_height=Cm(29.7)
    ns.left_margin=ns.right_margin=Cm(0.75); ns.top_margin=ns.bottom_margin=Cm(0.6)
    sp=ns._sectPr; cs=sp.xpath('./w:cols')[0]
    cs.set(qn('w:num'),str(numcols)); cs.set(qn('w:space'),'200')
    if numcols>1: cs.set(qn('w:sep'),'1')
    return ns

# ============ 頁首橫帶（全寬・開卷 30 秒用完）============
cols(1)
p=doc.add_paragraph(); tight(p,0,1,0.9); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('敷地計畫與都市設計 80150｜一頁作戰卡'); r.font.size=Pt(13); r.font.bold=True; r.font.color.rgb=NAVY; _f(r)
r2=p.add_run('　　申論30＋設計70｜4小時'); r2.font.size=Pt(8.5); r2.font.color.rgb=GRAY; _f(r2)

band('⏱ 0:00 申論 ▸ 0:50 審題＋抄圖說清單＋量體算式 ▸ 1:20 構想 ▸ 1:40 草圖 ▸ 2:10 正式圖 ▸ 3:40 標註＋檢核 ▸ 3:55 收',
     fill='1F3B5C', size=9.4, bold=True)
for c in doc.tables[-1].rows[0].cells:
    for pp in c.paragraphs:
        for rr in pp.runs: rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)

rich([('落後就砍：',True,RED),('分析圖細緻度→剖面小圖→透視縮小（可小不可無）　',False,None),
      ('絕不砍：',True,RED),('圖說清單每項・量體算式 ok!・剖面每層樓高・透視',False,None)],size=9.5,after=0.8)
rich([('版面　',True,NAVY),('左1/5 四格分析｜中1/2 全區配置圖＋剖立面｜右1/4 主標＋副標＋申論　',False,None),
      ('主標 4–6 字＋「/」串題目機能　',True,None),
      ('五色　',True,NAVY),('綠植栽・藍水體與引線文字・黃橘開放空間・紅重點與開挖線・灰周邊建物襯底',False,None)],size=9.5,after=2.0)

# ============ 兩欄 ============
newsec(2)

# ---------- 左欄：申論期 ----------
head('申論 30 分｜0:00–0:50　寫完不再回看')
rich([('起手・45字・照抄　',True,RED),('本案依都市計畫法主要計畫（§15）之指導、細部計畫（§22）之土地使用分區管制，並依都市設計準則管制量體與開放空間，→〔接本年議題〕',False,None)],after=0.5)
rich([('4個不套　',True,None),('純理論programming→不起手，只結尾橋接｜題目已括出法規面→豁免，直接破題｜都更vs危老→另一法系，不套§15/§32｜公設立體化→改套「多目標使用辦法」',False,None)],after=0.5)
rich([('收尾　',True,None),('全案於建築技術規則（停車／無障礙／防火避難／綠建築）基本盤上，導入〔依題旨選用：光電／綠屋頂／雨水回收／滯洪／韌性／文資保存〕，落實安全、永續而具公共性之實質環境。→「以下就設計說明之」',False,None)],after=0.5)
rich([('紀律　',True,RED),('400–600字＝全張唯一可成段處｜45–50分，',False,None),('過55分立刻停筆跳設計',True,None),('（70＞30）｜每點＝論點＋1句法源＋「見圖①」｜法源全篇≤5句',False,None)],after=0.5)
rich([('骨架反射・先畫框30秒　',True,None),('比較→兩欄表〔性質|範圍|機制|誘因|規模〕＋末列共同點｜角色功能→分支圖＋三欄表｜流程→箭頭圖｜基地分析→泡泡圖＋環境小剖面｜韌性減洪→要素表＋街道斷面',False,None)],after=0.5)
rich([('法源鏈　',True,None),('主計(§15)定位置｜細計(§22)定數字｜§32 分區得分級管制｜都審定形態｜技則定底線',False,None)],after=0.2)
para('└ 都審依都市計畫書、土管要點、都市設計準則及地方規範辦理（勿寫「無明文授權」這種法律結論）',indent=0.15,after=0.2)
para('└ 開挖限制屬「都市設計應表明事項」（通盤檢討辦法）→納入細部計畫／都市設計準則（非 §22 條文本身）',indent=0.15,after=0.2)
para('└「敷地計畫」為學科用語；法規對應概念＝建築基地／法定空地（建築法 §11）。勿寫「依敷地計畫法」',indent=0.15,after=1.2)

# ---------- 右欄（續排）：設計期 ----------
head('設計 70 分｜0:50–3:55')
rich([('0:00 先看試題「注意事項」　',True,RED),('計算器可否使用逐年不同，別憑上次經驗。',False,None)],after=0.3)
rich([('0:50 抄圖說清單→畫完打✓　',True,RED),('用詞照抄題目不改字｜缺1張＝該項拿不到完整分，小張也有分，空白＝0',False,None)],after=0.3)
para('□全區配置　□空間/機能組織（水平＋垂直）　□剖立面（要幾向畫幾向）　□外部空間局部透視（要幾處畫幾處）←114漏此張　□地下開挖範圍　□指北針・比例尺 S=1/___・圖名',indent=0.15,after=0.2)
para('※一樓平面＝看題目：113「含一樓平面」→必畫；108 明示「建築平面圖並非必要」→不畫。題目沒點名就不花時間',indent=0.15,after=0.5)

rich([('題目條件對照　',True,None),('圈題眼→每條寫一行→',False,None),('圖上真的指得出來才打✓',True,None),('（打不了✓＝還沒回應，不是還沒寫）',False,None)],after=0.2)
para('退縮○M｜老樹｜遺構水圳（→不開挖）｜高架噪音｜淹水潛勢｜汽機車出入口與人行分設（兩支箭頭各自標）',indent=0.15,after=0.5)

rich([('放樣・第一筆　',True,RED),('不是量體，是開放空間；起點釘在英雄條件（老樹／水岸／捷運口／古蹟／眺望軸）←治 T9。量體讓開放空間（虛空是正空間，不是剩料）。泡泡圖一句：以○為核心，串聯○－○－○',False,None)],after=0.5)

rich([('量體對帳・三行　',True,None),('先算上限→上限÷層數＝每層可用→照此畫→最後回抄實際值（',False,None),('勿畫完再湊',True,RED),('）',False,None)],after=0.2)
para('①建面＝基地×建蔽　②法容＝基地×容積〔上限〕　③本案＝逐棟樓地板相加　→ 收尾必寫 ③≤② …ok!',indent=0.15,after=0.2)
para('不合→改方塊，絕不改②｜題目給「坪」先×3.3｜法定停車空間「得不計入」容積（技則§162，有但書＋地方規定）｜1車位≈40㎡',indent=0.15,after=0.5)

rich([('引線標籤　',True,None),('每則5–15字，超過就砍｜配置8–12則（藍筆環繞一圈）｜剖面5–8｜分析圖每張1',False,None)],after=0.2)
para('①名詞－效益：退縮建築－減少壓迫感　②作為＋數據＋簡稱：開挖率65%<70%｜保水(技則)',indent=0.15,after=0.2)
rich([('③英雄扣回×3則以上（治T9）：',False,None),('〔開放空間〕順應〔老樹／水岸／捷運口〕→使…',True,None)],indent=0.15,after=0.2)
para('✗只寫名詞　✗超過15字　✗圖上寫條號（一律簡稱）',indent=0.15,after=0.5)

rich([('剖面・三件事缺一就失分　',True,RED),('畫大：長≥配置圖寬、高≥8cm；剖線切過開放空間＋主量體',False,None)],after=0.2)
para('①頂＝剖到每段名稱＋寬　②右緣＝每層樓高數字（114敗因）　③紅[分類]×5：室內環境／日常節能／基地保水／水資源／廢棄物減量',indent=0.15,after=0.2)
para('＋地下層數・覆土≥1.0m・防水 FL＋○.○○・天際線｜淨高 羽球6.1／球場7–9／停車2.4／騎樓3.4｜層高 一般3.6–4.2／大廳4.5–6／地下停車3.5',indent=0.15,after=0.5)

rich([('透視・題目要求就必畫　',True,RED),('開卷即鉛筆框 15×10cm，末30分填滿。醜不扣分，缺了該項拿不到分',False,None)],after=0.2)
para('下筆3問全YES：①英雄條件入鏡　②有人在動　③看得到量體關係',indent=0.15,after=0.2)
para('配字2–3行＝〔場所名〕－〔角色〕｜〔誰〕〔做什麼〕　例：深廊－串聯廣場與圖書館的中介空間；午後長者下棋、學童候車',indent=0.15,after=0.2)
para('人物≥8＋喬木＋對面建物淡灰襯底｜框內不寫法規、不寫尺寸',indent=0.15,after=0.5)

rich([('手要寫字→改畫這個　',True,RED),('申論區外任一處滿3行→停筆',False,None)],after=0.2)
para('層級→泡泡圖｜動線→箭頭＋圖例｜垂直機能→寫進量體格（RF屋頂花園／5F羽球場）｜因果→名詞－效益｜計算→算式＋ok!｜比較檢核→兩欄表｜多點→①②③↔圖上同編號｜次階用「└」',indent=0.15,after=0.5)

rich([('水綠4數字　',True,None),('滯洪 ≥基地×0.045 m³/㎡〔技則§4-3：都計區新增改建，300㎡以下等除外；地方另有規定從其規定〕｜保水 >0.5×法定空地比率〔§305，適用範圍另見§298〕｜綠化用喬木複層非草皮｜出流管制 ≥1ha・都計區新建 ≥0.2ha〔現行新制〕',False,None)],after=0.5)

rich([('法規檢核表　',True,None),('配置圖角落 5×8cm：項目｜法定｜本案｜✓　列：建蔽・容積・高度・退縮・法停・開挖率・綠覆率。無法定值者填「—」照列＝主動控制',False,None)],after=0.5)

rich([('開挖率／開挖範圍　',True,None),('近年連三年考（112/113/114）。配置圖＋一樓平面用',False,None),('粗虛線',True,None),('框地下室外緣，標「地下開挖範圍」',False,None)],after=0.2)
para('算式寫圖旁：開挖面積÷基地面積＝○%　上限依題目／土管／都審；無題示則自控 60–70% 並註明「自訂控制值」',indent=0.15,after=0.2)
para('虛線須避開：退縮帶／老樹樹冠根系／遺構水圳／鄰地界線',indent=0.15,after=0.2)
para('旁註 覆土≥1.0m〔設計控制值，依地方都審／植栽規範〕・防水高程 FL＋○.○○',indent=0.15,after=0.5)

rich([('模組庫・卡住就調用　',True,None),('左欄四格（基地分析理念／動線防救災／開放空間／彈性格）｜基地環境小剖面①–⑥｜Lynch五元素＋5小圖示｜防災8–9點＋等角圖｜生態工法細部剖面｜開放空間泡泡圖｜街道斷面',False,None)],after=0.8)

# ============ 頁尾（全寬）============
newsec(1)
rich([('※ 本卡的 ',False,GRAY),('字數/則數/尺寸（400–600字・8–12則・15×10cm・人物≥8・剖面高≥8cm）',True,GRAY),
      (' 為自訂時間與版面控制值，非考選部評分標準；法規數字請以全國法規資料庫與題目給定為準。',False,GRAY)],size=7.6,after=1.0)
band('交卷前 90 秒　□題目點名的圖全到齊（透視！）　□圖說清單逐項✓　□每條算式數字對得上＋ok!　□剖面每層有樓高數字　□透視有人物　□指北針・比例尺・圖名　□題目條件欄全✓',
     fill='C00000', size=9.4, bold=True)
for c in doc.tables[-1].rows[0].cells:
    for pp in c.paragraphs:
        for rr in pp.runs: rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)

out='/home/user/architecture-study-notes/講義/敷地A4作戰卡.docx'
doc.save(out); print('saved:',out)
